"""
Generates monthly reports in both Word (.docx) and PDF formats.
"""

from __future__ import annotations

import calendar
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF
from loguru import logger

from config.settings import REPORTS_DIR, REPORT_AUTHOR, REPORT_COMPANY


class ReportGenerator:
    def generate(
        self,
        logs: list[dict],
        trend_summary: str,
        year: int,
        month: int,
    ) -> dict[str, Path]:
        month_name = calendar.month_name[month]
        title = f"Tech Market Report — {month_name} {year}"

        docx_path = self._generate_docx(title, logs, trend_summary, year, month)
        pdf_path = self._generate_pdf(title, logs, trend_summary, year, month)

        logger.info(f"Report saved → {docx_path} | {pdf_path}")
        return {"docx": docx_path, "pdf": pdf_path}

    # ── Word ───────────────────────────────────────────────────────────────────

    def _generate_docx(
        self,
        title: str,
        logs: list[dict],
        trend_summary: str,
        year: int,
        month: int,
    ) -> Path:
        doc = Document()

        # Title
        heading = doc.add_heading(title, level=0)
        heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        meta = f"Generated: {datetime.now().strftime('%Y-%m-%d')}  |  Author: {REPORT_AUTHOR}"
        if REPORT_COMPANY:
            meta += f"  |  {REPORT_COMPANY}"
        doc.add_paragraph(meta).runs[0].font.size = Pt(9)

        doc.add_heading("Monthly Trend Analysis", level=1)
        doc.add_paragraph(trend_summary)

        # Group by source
        by_source: dict[str, list[dict]] = {}
        for item in logs:
            by_source.setdefault(item.get("source", "Unknown"), []).append(item)

        doc.add_heading("Collected Items by Source", level=1)
        for source, items in sorted(by_source.items()):
            doc.add_heading(source, level=2)
            for item in items[:20]:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(item.get("title", ""))
                run.bold = True
                summary = item.get("llm_summary") or item.get("summary", "")
                if summary:
                    p.add_run(f"\n{summary[:300]}")
                url = item.get("url", "")
                if url:
                    p.add_run(f"\n{url}").font.size = Pt(8)

        filename = f"report_{year}_{month:02d}.docx"
        path = REPORTS_DIR / filename
        doc.save(str(path))
        return path

    # ── PDF ────────────────────────────────────────────────────────────────────

    def _generate_pdf(
        self,
        title: str,
        logs: list[dict],
        trend_summary: str,
        year: int,
        month: int,
    ) -> Path:
        pdf = _TechReportPDF(title=title)
        pdf.add_page()

        pdf.section_title("Monthly Trend Analysis")
        pdf.body_text(trend_summary)

        by_source: dict[str, list[dict]] = {}
        for item in logs:
            by_source.setdefault(item.get("source", "Unknown"), []).append(item)

        pdf.section_title("Collected Items by Source")
        for source, items in sorted(by_source.items()):
            pdf.subsection_title(source)
            for item in items[:20]:
                item_title = item.get("title", "")
                summary = (item.get("llm_summary") or item.get("summary", ""))[:200]
                pdf.bullet_item(item_title, summary)

        filename = f"report_{year}_{month:02d}.pdf"
        path = REPORTS_DIR / filename
        pdf.output(str(path))
        return path


class _TechReportPDF(FPDF):
    def __init__(self, title: str) -> None:
        super().__init__()
        self._report_title = title

    def header(self) -> None:
        self.set_font("Helvetica", "B", 10)
        self.set_text_color(26, 86, 219)
        self.cell(0, 8, self._report_title, align="L", new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(26, 86, 219)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def footer(self) -> None:
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 8, f"Page {self.page_no()}", align="C")

    def section_title(self, text: str) -> None:
        self.ln(4)
        self.set_font("Helvetica", "B", 13)
        self.set_text_color(26, 86, 219)
        self.multi_cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)
        self.set_text_color(0, 0, 0)

    def subsection_title(self, text: str) -> None:
        self.ln(3)
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(50, 50, 50)
        self.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
        self.set_text_color(0, 0, 0)

    def body_text(self, text: str) -> None:
        self.set_font("Helvetica", "", 10)
        self.multi_cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def bullet_item(self, title: str, summary: str) -> None:
        self.set_font("Helvetica", "B", 9)
        self.multi_cell(0, 5, f"• {title}", new_x="LMARGIN", new_y="NEXT")
        if summary:
            self.set_font("Helvetica", "", 8)
            self.set_x(self.l_margin + 4)
            self.multi_cell(0, 4, summary, new_x="LMARGIN", new_y="NEXT")
        self.ln(1)
