from __future__ import annotations

import json
import logging
import os
import re
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv
from openai import OpenAI

from src.config import PROJECT_ROOT, Settings

logger = logging.getLogger(__name__)

CATEGORY_LABELS = {
    "tech_news": "Tech News",
    "academic": "Academic",
    "enterprise": "Enterprise",
}

CATEGORY_LABELS_KO = {
    "tech_news": "테크 뉴스",
    "academic": "학술",
    "enterprise": "기업",
}

_HEADER_BLUE = RGBColor(0x1F, 0x39, 0x64)
_SUBTITLE_BLUE = RGBColor(0x44, 0x54, 0x6A)

# ---------------------------------------------------------------------------
# JSON prompt templates
# ---------------------------------------------------------------------------

_SCHEMA_JSON = """{
  "technology_name": "Main technology theme (e.g. 'Generative AI & LLMs')",
  "monthly_headline": "1-2 sentence headline capturing the single most important development this month. Must include [N] citation.",
  "monthly_context": "3-4 sentences explaining the structural shift or opportunity the month's articles collectively point to. Include specific figures and company names. End each sentence with [N].",
  "sec1": {
    "snapshot": "3-sentence plain-language description. Lead with so-what. End each sentence with [N].",
    "key_findings": [
      "Market signal: ... [N]",
      "Competitive signal: ... [N]",
      "Korea-specific signal: ... [N]",
      "Risk signal: ... [N]"
    ],
    "metrics": [
      {"metric": "Global Market Size",        "value": "Only if explicitly stated in articles, else N/A", "yoy": "Only if stated, else –", "forecast": "Only if stated, else N/A", "source": "Exact source name from the article that provides this figure. Never invent."},
      {"metric": "Korea Market Size",          "value": "Only if explicitly stated in articles, else N/A", "yoy": "Only if stated, else –", "forecast": "Only if stated, else N/A", "source": "Exact source name from the article that provides this figure. Never invent."},
      {"metric": "TRL Level",                  "value": "Only if stated in articles, else N/A", "yoy": "–", "forecast": "Only if stated, else N/A", "source": "Exact source name. Never invent."},
      {"metric": "Top Filing Country (Patents)","value": "Only if stated in articles, else N/A", "yoy": "–", "forecast": "–", "source": "Exact source name. Never invent."},
      {"metric": "Leading Vendor Market Share", "value": "Only if stated in articles, else N/A", "yoy": "–", "forecast": "–", "source": "Exact source name. Never invent."}
    ]
  },
  "sec2": {
    "definition": "Core function, underlying principles, key variants. [N]",
    "trl_table": [
      {"dimension": "Current TRL",                  "assessment": "TRL X",  "basis": "Standard/Fraunhofer internal assessment"},
      {"dimension": "Target TRL",                   "assessment": "TRL X",  "basis": "Product/project roadmap"},
      {"dimension": "Estimated time to TRL 9",      "assessment": "X years","basis": "Expert assessment/foresight study"},
      {"dimension": "Comparable technology benchmark","assessment": "Tech", "basis": "Benchmark report"}
    ],
    "differentiation": "Performance envelope, cost structure, IP position, maturity, Korea adoption trajectory. [N]",
    "alt_a_name": "Alternative A name",
    "alt_b_name": "Alternative B name",
    "comparison_table": [
      {"feature": "Performance",    "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "Cost",           "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "Maturity",       "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "IP Position",    "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "Korea Adoption", "this_tech": "...", "alt_a": "...", "alt_b": "..."}
    ],
    "patents": [
      {"no": "EP-XXXXXXX", "title": "Title", "assignee": "Company", "country": "Country", "year": "Year", "source": "Espacenet"},
      {"no": "KR-XXXXXXX", "title": "Title", "assignee": "Company", "country": "Korea",   "year": "Year", "source": "KIPO"}
    ]
  },
  "sec3": {
    "overview": "Global market size ($XB), CAGR, 5-year trajectory, forecast horizon. [N]",
    "segmentation": [
      {"type": "By Application", "name": "App1/App2",          "size": "$XB", "share": "X%", "growth": "X% CAGR", "notes": "Note"},
      {"type": "By End-User",    "name": "Industry sector",    "size": "$XB", "share": "X%", "growth": "X% CAGR", "notes": "Note"},
      {"type": "By Deployment",  "name": "Deployment model",   "size": "$XB", "share": "X%", "growth": "X% CAGR", "notes": "Note"},
      {"type": "By Geography",   "name": "Korea/APAC/EU/US",   "size": "$XB", "share": "X%", "growth": "X% CAGR", "notes": "Note"}
    ],
    "regional": [
      {"region": "Korea",  "size": "$XB", "drivers": "Driver", "policy": "MOTIE, MSIT, IITP",    "players": "Company/Institute", "source": "KIET/KOTRA"},
      {"region": "Japan",  "size": "$XB", "drivers": "Driver", "policy": "METI, NEDO",            "players": "Company",           "source": "METI"},
      {"region": "China",  "size": "$XB", "drivers": "Driver", "policy": "MIIT, CAICT",           "players": "Company",           "source": "CAICT"},
      {"region": "EU",     "size": "$XB", "drivers": "Driver", "policy": "EU Commission, BMBF",   "players": "Company",           "source": "Eurostat"},
      {"region": "US",     "size": "$XB", "drivers": "Driver", "policy": "NIST, DOE",             "players": "Company",           "source": "IDC"},
      {"region": "SE Asia","size": "$XB", "drivers": "Driver", "policy": "EDB Singapore, IMDA",   "players": "Company",           "source": "ADB"},
      {"region": "India",  "size": "$XB", "drivers": "Driver", "policy": "NITI Aayog, DST",       "players": "Company",           "source": "ADB"}
    ],
    "drivers_barriers": [
      {"driver": "Driver 1 [N]", "barrier": "Barrier 1 [N]"},
      {"driver": "Driver 2 [N]", "barrier": "Barrier 2 [N]"},
      {"driver": "Driver 3 [N]", "barrier": "Barrier 3 [N]"}
    ]
  },
  "sec4": {
    "vendors": [
      {"vendor": "Global Vendor A", "hq": "Country", "type": "Private",       "offering": "Product/service", "strategy": "Strategy", "position": "Leader/Challenger",   "source": "IDC MarketScape"},
      {"vendor": "Global Vendor B", "hq": "Country", "type": "Public",        "offering": "Product/service", "strategy": "Strategy", "position": "Niche/Contender",     "source": "Gartner MQ"},
      {"vendor": "ETRI / KAIST",    "hq": "Korea",   "type": "Public Research","offering": "R&D programme",  "strategy": "Gov-backed applied research", "position": "Key domestic player", "source": "KOTRA/KIET"}
    ],
    "korea_context": "Government-backed players (ETRI, KAIST, POSTECH), chaebol, startup ecosystem. [N]",
    "swot": {
      "strengths":     ["Strength 1 [N]",     "Strength 2",     "Strength 3"],
      "weaknesses":    ["Weakness 1 [N]",     "Weakness 2",     "Weakness 3"],
      "opportunities": ["Opportunity 1 [N]",  "Opportunity 2",  "Opportunity 3"],
      "threats":       ["Threat 1 [N]",       "Threat 2",       "Threat 3"]
    },
    "five_forces": [
      {"force": "Supplier Power", "intensity": "High/Med/Low", "key_factor": "Factor", "implication": "Implication"},
      {"force": "Buyer Power",    "intensity": "High/Med/Low", "key_factor": "Factor", "implication": "Implication"},
      {"force": "Rivalry",        "intensity": "High/Med/Low", "key_factor": "Factor", "implication": "Implication"},
      {"force": "New Entrants",   "intensity": "High/Med/Low", "key_factor": "Factor", "implication": "Implication"},
      {"force": "Substitutes",    "intensity": "High/Med/Low", "key_factor": "Factor", "implication": "Implication"}
    ]
  },
  "sec5": {
    "publications": [
      {"title": "Paper title [N]", "authors": "Author A et al.", "journal": "Journal/Conference", "year": "Year", "citations": "X", "doi": "doi:XX.XXXX/XXXXX"},
      {"title": "Paper title [N]", "authors": "Author B et al.", "journal": "Journal/Conference", "year": "Year", "citations": "X", "doi": "doi:XX.XXXX/XXXXX"}
    ],
    "funding": [
      {"program": "Program name", "body": "EU Horizon Europe", "region": "EU",      "budget": "€XB",  "focus": "Focus", "timeline": "Years", "source": "ec.europa.eu"},
      {"program": "Program name", "body": "BMBF",              "region": "Germany", "budget": "€XM",  "focus": "Focus", "timeline": "Years", "source": "bmbf.de"},
      {"program": "Program name", "body": "IITP",              "region": "Korea",   "budget": "₩XB",  "focus": "Focus", "timeline": "Years", "source": "iitp.kr"},
      {"program": "Program name", "body": "KISTEP",            "region": "Korea",   "budget": "₩XB",  "focus": "Focus", "timeline": "Years", "source": "kistep.re.kr"},
      {"program": "Program name", "body": "NEDO",              "region": "Japan",   "budget": "¥XB",  "focus": "Focus", "timeline": "Years", "source": "nedo.go.jp"},
      {"program": "Program name", "body": "NIST",              "region": "US",      "budget": "$XM",  "focus": "Focus", "timeline": "Years", "source": "nist.gov"}
    ],
    "emerging_directions": [
      "Direction 1: emerging research area at intersection with AI/materials science. [N]",
      "Direction 2: convergence trend or cross-sector application gaining traction. [N]",
      "Direction 3: white-space R&D opportunity relevant to Fraunhofer Korea's mandate. [N]"
    ]
  },
  "sec6": {
    "policies": [
      {"policy": "EU Act/Directive", "body": "EU Commission", "region": "EU",     "status": "Active/Pending", "date": "Date", "impact": "High/Med/Low", "source": "ec.europa.eu"},
      {"policy": "Korean policy",    "body": "MOTIE/MSIT",    "region": "Korea",  "status": "Active/Pending", "date": "Date", "impact": "High/Med/Low", "source": "motie.go.kr"},
      {"policy": "Japanese policy",  "body": "METI",          "region": "Japan",  "status": "Active/Pending", "date": "Date", "impact": "High/Med/Low", "source": "meti.go.jp"},
      {"policy": "US standard",      "body": "NIST",          "region": "US",     "status": "Active/Pending", "date": "Date", "impact": "High/Med/Low", "source": "nist.gov"},
      {"policy": "Intl standard",    "body": "OECD/ISO",      "region": "Global", "status": "Active/Pending", "date": "Date", "impact": "High/Med/Low", "source": "oecd.org"}
    ],
    "compliance": [
      "Data governance: applicable data localisation/privacy requirements. [N]",
      "Intellectual property: IP licensing obligations, SEP considerations. [N]",
      "Export control: dual-use classifications, Wassenaar items. [N]"
    ]
  },
  "sec7": {
    "hype_cycle_phase": "Innovation Trigger / Peak of Inflated Expectations / Trough of Disillusionment / Slope of Enlightenment / Plateau of Productivity",
    "hype_cycle_plateau": "X years",
    "hype_cycle_source": "Gartner [Year] Hype Cycle for [Category]",
    "predictions": [
      {"year": "Year+1", "prediction": "Prediction [N]", "confidence": "High/Med/Low", "implication": "Implication", "source": "IDC"},
      {"year": "Year+2", "prediction": "Prediction",     "confidence": "High/Med/Low", "implication": "Implication", "source": "IDC"},
      {"year": "Year+3", "prediction": "Prediction",     "confidence": "High/Med/Low", "implication": "Implication", "source": "IDC/Gartner"},
      {"year": "Year+4", "prediction": "Prediction",     "confidence": "High/Med/Low", "implication": "Implication", "source": "Gartner"},
      {"year": "Year+5", "prediction": "Prediction",     "confidence": "High/Med/Low", "implication": "Implication", "source": "McKinsey MGI"}
    ],
    "roadmap": [
      {"phase": "Short-term", "timeframe": "0–2 years",  "milestone": "Milestone [N]", "trl": "TRL X", "enabler": "Enabler", "risk": "Risk"},
      {"phase": "Mid-term",   "timeframe": "2–5 years",  "milestone": "Milestone",     "trl": "TRL X", "enabler": "Enabler", "risk": "Risk"},
      {"phase": "Long-term",  "timeframe": "5–10 years", "milestone": "Milestone",     "trl": "TRL X", "enabler": "Enabler", "risk": "Risk"}
    ]
  },
  "sec8": {
    "opportunities": [
      {"opportunity": "Joint R&D",          "type": "Research",         "partners": "ETRI/KAIST/POSTECH",  "funding": "IITP/BMBF",       "priority": "High/Med/Low", "timeline": "Year"},
      {"opportunity": "Tech Transfer",      "type": "Commercialization","partners": "Korean SME/Chaebol",  "funding": "MOTIE/KISTEP",    "priority": "High/Med/Low", "timeline": "Year"},
      {"opportunity": "Policy Contribution","type": "Advisory",         "partners": "MSIT/MOTIE",          "funding": "–",               "priority": "High/Med/Low", "timeline": "Year"},
      {"opportunity": "EU–Korea Bridge",    "type": "Collaboration",    "partners": "EU research partner", "funding": "Horizon Europe",  "priority": "High/Med/Low", "timeline": "Year"}
    ],
    "risks": [
      {"risk": "Risk 1 [N]", "likelihood": "H/M/L", "impact": "H/M/L", "mitigation": "Mitigation", "owner": "Function"},
      {"risk": "Risk 2",     "likelihood": "H/M/L", "impact": "H/M/L", "mitigation": "Mitigation", "owner": "Function"},
      {"risk": "Risk 3",     "likelihood": "H/M/L", "impact": "H/M/L", "mitigation": "Mitigation", "owner": "Function"}
    ],
    "actions": [
      "Action 1: Schedule structured technology briefing with ETRI/KISTEP. [N]",
      "Action 2: Initiate patent freedom-to-operate (FTO) analysis via Espacenet and KIPO.",
      "Action 3: Draft Horizon Europe partnership proposal with Korean research university or chaebol R&D lab."
    ]
  },
  "sec9": {
    "methodology": "Describe automated data collection: daily API pulls, keyword list and Boolean search strings, deduplication and relevance-scoring logic, monthly report generation process and QA.",
    "quality": "Identify data gaps by source type; state confidence levels (High/Med/Low) for each major metric; document update frequency per source."
  }
}"""

_PROMPT_EN = """You are a technology market research analyst at Fraunhofer Institute Korea Office.
Analyse the {count} monitoring articles from {year}-{month:02d} below and produce a structured JSON report.

CITATION RULES (strictly enforced – do not skip):
• Every factual claim, statistic, or trend statement MUST end with [N] where N is the article 'ref' number.
• Use multiple citations [1][2] when relevant.
• Do NOT invent or extrapolate facts absent from the provided data.
• If a field cannot be substantiated, write "N/A".

METRICS DASHBOARD RULES (never violate):
• Each metric's "value", "yoy", and "forecast" must only be filled when the exact figure appears in the provided articles.
  If not present: value → "N/A", yoy → "–", forecast → "N/A".
• The "source" field must be the exact source name (e.g. "IEA", "Financial Times Tech") of the article containing that figure.
  Never use generic labels like "Gartner/IDC" or "KIET/KOTRA" unless those exact names appear in the articles.
• TRL, patent filing country, and vendor market share must all be "N/A" unless explicitly stated in an article.

monthly_headline / monthly_context RULES:
• monthly_headline: The single most important event or trend this month in 1-2 sentences with specific figures.
• monthly_context: 3-4 sentences that let a reader grasp this month's significance without reading further.

sec1.key_findings RULES (strictly enforced):
• MUST have exactly 4 items: Market signal, Competitive signal, Korea-specific signal, Risk signal.
• Each item must cover ALL domains in the articles (semiconductors, energy, enterprise, etc.) – do not focus on just one sector.
• Market signal: Global market size, investment, or demand facts with specific figures and company names.
• Competitive signal: A specific company's strategic move (capacity expansion, market capture, M&A, etc.).
• Korea-specific signal: Content directly relevant to Korea's market, policy, or companies.
• Risk signal: Supply chain bottlenecks, regulatory barriers, or technology constraints.
• Each finding must be a complete, specific sentence with figures and citations – never abbreviate or omit key facts.

sec2.definition RULES:
• The first sentence MUST explicitly name the technology using the exact string from the technology_name field.
• Do not write "This technology..." or vague references – name it directly.

N/A RULES:
• If a text field (overview, definition, etc.) cannot be substantiated, set the entire field to "N/A" – never write a sentence that contains N/A as a value.
  Wrong: "The global market size of AI-powered grid management is N/A [N/A]."
  Correct: "N/A"
• In trl_table, if the actual TRL level (1–9) is unknown, set assessment to "N/A". Never leave template placeholders like "TRL X" or "X years".

OUTPUT: Return ONLY valid JSON (no markdown fences, no commentary) matching this schema exactly:
{schema}

Indexed article data (JSON):
{articles_json}"""

_SCHEMA_JSON_KO = """{
  "technology_name": "주요 기술 테마 (예: '생성 AI & LLM')",
  "monthly_headline": "이달 가장 중요한 한 가지 발전을 1~2문장으로 서술. 반드시 [N] 인용 포함.",
  "monthly_context": "월간 핵심 맥락: 이달 기사들이 공통적으로 가리키는 구조적 변화 또는 기회를 3~4문장으로 서술. 수치·기업명 포함, 각 문장 끝에 [N] 인용.",
  "sec1": {
    "snapshot": "3문장 기술 개요. '무엇이 중요한가'를 먼저 서술. 각 문장 끝에 [N] 인용.",
    "key_findings": [
      "시장 신호: ... [N]",
      "경쟁 신호: ... [N]",
      "한국 특화 신호: ... [N]",
      "리스크 신호: ... [N]"
    ],
    "metrics": [
      {"metric": "글로벌 시장 규모",      "value": "기사에 명시된 수치만 기입, 없으면 N/A", "yoy": "기사에 명시된 경우만, 없으면 –", "forecast": "기사에 명시된 경우만, 없으면 N/A", "source": "반드시 기사의 실제 출처명 그대로 (예: IEA, Financial Times Tech 등). 추정 금지."},
      {"metric": "한국 시장 규모",        "value": "기사에 명시된 수치만 기입, 없으면 N/A", "yoy": "기사에 명시된 경우만, 없으면 –", "forecast": "기사에 명시된 경우만, 없으면 N/A", "source": "반드시 기사의 실제 출처명 그대로. 추정 금지."},
      {"metric": "TRL 단계",             "value": "기사에 명시된 경우만 기입, 없으면 N/A", "yoy": "–", "forecast": "기사에 명시된 경우만, 없으면 N/A", "source": "반드시 기사의 실제 출처명 그대로. 추정 금지."},
      {"metric": "주요 특허 출원국",      "value": "기사에 명시된 경우만 기입, 없으면 N/A", "yoy": "–", "forecast": "–", "source": "반드시 기사의 실제 출처명 그대로. 추정 금지."},
      {"metric": "선도 벤더 시장점유율",  "value": "기사에 명시된 경우만 기입, 없으면 N/A", "yoy": "–", "forecast": "–", "source": "반드시 기사의 실제 출처명 그대로. 추정 금지."}
    ]
  },
  "sec2": {
    "definition": "핵심 기능, 작동 원리, 주요 변형 기술 설명. [N]",
    "trl_table": [
      {"dimension": "현재 TRL",              "assessment": "TRL X",  "basis": "기준/프라운호퍼 내부 평가"},
      {"dimension": "목표 TRL",              "assessment": "TRL X",  "basis": "제품/프로젝트 로드맵"},
      {"dimension": "TRL 9 도달 예상 기간",  "assessment": "X년",    "basis": "전문가 평가/미래예측 연구"},
      {"dimension": "비교 기술 벤치마크",    "assessment": "기술명",  "basis": "출처/벤치마크 보고서"}
    ],
    "differentiation": "성능 범위, 비용 구조, 지재권 위치, 성숙도, 한국 내 도입 현황. [N]",
    "alt_a_name": "대안 기술 A 이름",
    "alt_b_name": "대안 기술 B 이름",
    "comparison_table": [
      {"feature": "성능",         "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "비용",         "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "성숙도",       "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "IP 위치",      "this_tech": "...", "alt_a": "...", "alt_b": "..."},
      {"feature": "한국 도입 현황","this_tech": "...", "alt_a": "...", "alt_b": "..."}
    ],
    "patents": [
      {"no": "EP-XXXXXXX", "title": "특허 제목", "assignee": "출원인명", "country": "국가", "year": "연도", "source": "Espacenet"},
      {"no": "KR-XXXXXXX", "title": "특허 제목", "assignee": "출원인명", "country": "한국", "year": "연도", "source": "KIPO"}
    ]
  },
  "sec3": {
    "overview": "글로벌 시장 규모($X억달러), CAGR(연평균 성장률), 과거 5년 추이, 예측 기간 제공. [N]",
    "segmentation": [
      {"type": "응용 분야별",  "name": "앱1/앱2",           "size": "$X억달러", "share": "X%", "growth": "X% CAGR", "notes": "비고"},
      {"type": "최종 사용자별","name": "산업 분야",          "size": "$X억달러", "share": "X%", "growth": "X% CAGR", "notes": "비고"},
      {"type": "배포 방식별",  "name": "배포 모델",          "size": "$X억달러", "share": "X%", "growth": "X% CAGR", "notes": "비고"},
      {"type": "지역별",       "name": "한국/APAC/EU/미국", "size": "$X억달러", "share": "X%", "growth": "X% CAGR", "notes": "비고"}
    ],
    "regional": [
      {"region": "한국",  "size": "$X억달러", "drivers": "동인", "policy": "MOTIE, MSIT, IITP",  "players": "기업/연구기관", "source": "KIET/KOTRA"},
      {"region": "일본",  "size": "$X억달러", "drivers": "동인", "policy": "METI, NEDO",          "players": "기업",         "source": "METI"},
      {"region": "중국",  "size": "$X억달러", "drivers": "동인", "policy": "MIIT, CAICT",         "players": "기업",         "source": "CAICT"},
      {"region": "EU",    "size": "$X억달러", "drivers": "동인", "policy": "EU Commission, BMBF", "players": "기업",         "source": "Eurostat"},
      {"region": "미국",  "size": "$X억달러", "drivers": "동인", "policy": "NIST, DOE",           "players": "기업",         "source": "IDC"},
      {"region": "동남아","size": "$X억달러", "drivers": "동인", "policy": "EDB Singapore, IMDA", "players": "기업",         "source": "ADB"},
      {"region": "인도",  "size": "$X억달러", "drivers": "동인", "policy": "NITI Aayog, DST",     "players": "기업",         "source": "ADB"}
    ],
    "drivers_barriers": [
      {"driver": "동인 1 [N]", "barrier": "장벽 1 [N]"},
      {"driver": "동인 2 [N]", "barrier": "장벽 2 [N]"},
      {"driver": "동인 3 [N]", "barrier": "장벽 3 [N]"}
    ]
  },
  "sec4": {
    "vendors": [
      {"vendor": "글로벌 벤더 A", "hq": "국가", "type": "민간",       "offering": "제품/서비스", "strategy": "전략", "position": "리더/챌린저",     "source": "IDC MarketScape"},
      {"vendor": "글로벌 벤더 B", "hq": "국가", "type": "상장",       "offering": "제품/서비스", "strategy": "전략", "position": "니치/경쟁자",     "source": "Gartner MQ"},
      {"vendor": "ETRI / KAIST",  "hq": "한국", "type": "공공 연구기관","offering": "R&D 프로그램","strategy": "정부 지원 응용 연구", "position": "핵심 국내 플레이어", "source": "KOTRA/KIET"}
    ],
    "korea_context": "한국 경쟁 환경: 정부 지원 기관(ETRI, KAIST, POSTECH), 재벌 참여, 스타트업 생태계. [N]",
    "swot": {
      "strengths":     ["강점 1 [N]", "강점 2", "강점 3"],
      "weaknesses":    ["약점 1 [N]", "약점 2", "약점 3"],
      "opportunities": ["기회 1 [N]", "기회 2", "기회 3"],
      "threats":       ["위협 1 [N]", "위협 2", "위협 3"]
    },
    "five_forces": [
      {"force": "공급자 교섭력", "intensity": "높음/중간/낮음", "key_factor": "요인", "implication": "시사점"},
      {"force": "구매자 교섭력", "intensity": "높음/중간/낮음", "key_factor": "요인", "implication": "시사점"},
      {"force": "경쟁 강도",    "intensity": "높음/중간/낮음", "key_factor": "요인", "implication": "시사점"},
      {"force": "신규 진입자",  "intensity": "높음/중간/낮음", "key_factor": "요인", "implication": "시사점"},
      {"force": "대체재",       "intensity": "높음/중간/낮음", "key_factor": "요인", "implication": "시사점"}
    ]
  },
  "sec5": {
    "publications": [
      {"title": "논문 제목 [N]", "authors": "저자 A 외", "journal": "저널/학회명", "year": "연도", "citations": "X", "doi": "doi:XX.XXXX/XXXXX"},
      {"title": "논문 제목 [N]", "authors": "저자 B 외", "journal": "저널/학회명", "year": "연도", "citations": "X", "doi": "doi:XX.XXXX/XXXXX"}
    ],
    "funding": [
      {"program": "프로그램명", "body": "EU Horizon Europe", "region": "EU",   "budget": "€X억",  "focus": "분야", "timeline": "기간", "source": "ec.europa.eu"},
      {"program": "프로그램명", "body": "BMBF",             "region": "독일",  "budget": "€X백만","focus": "분야", "timeline": "기간", "source": "bmbf.de"},
      {"program": "프로그램명", "body": "IITP",             "region": "한국",  "budget": "₩X억", "focus": "분야", "timeline": "기간", "source": "iitp.kr"},
      {"program": "프로그램명", "body": "KISTEP",           "region": "한국",  "budget": "₩X억", "focus": "분야", "timeline": "기간", "source": "kistep.re.kr"},
      {"program": "프로그램명", "body": "NEDO",             "region": "일본",  "budget": "¥X억", "focus": "분야", "timeline": "기간", "source": "nedo.go.jp"},
      {"program": "프로그램명", "body": "NIST",             "region": "미국",  "budget": "$X백만","focus": "분야", "timeline": "기간", "source": "nist.gov"}
    ],
    "emerging_directions": [
      "미래 연구 주제 1: AI·소재 과학 등과 교차점의 새로운 연구 영역. [N]",
      "미래 연구 주제 2: 학술 문헌에서 주목받는 기술 융합 트렌드. [N]",
      "미래 연구 주제 3: 프라운호퍼 한국의 응용 연구 사명 관련 공백 R&D 기회. [N]"
    ]
  },
  "sec6": {
    "policies": [
      {"policy": "EU 법안/지침명",  "body": "EU Commission", "region": "EU",    "status": "시행/심의", "date": "날짜", "impact": "높음/중간/낮음", "source": "ec.europa.eu"},
      {"policy": "한국 정책명",     "body": "MOTIE/MSIT",    "region": "한국",  "status": "시행/심의", "date": "날짜", "impact": "높음/중간/낮음", "source": "motie.go.kr"},
      {"policy": "일본 정책명",     "body": "METI",          "region": "일본",  "status": "시행/심의", "date": "날짜", "impact": "높음/중간/낮음", "source": "meti.go.jp"},
      {"policy": "미국 표준/정책",  "body": "NIST",          "region": "미국",  "status": "시행/심의", "date": "날짜", "impact": "높음/중간/낮음", "source": "nist.gov"},
      {"policy": "국제 표준",       "body": "OECD/ISO",      "region": "글로벌","status": "시행/심의", "date": "날짜", "impact": "높음/중간/낮음", "source": "oecd.org"}
    ],
    "compliance": [
      "데이터 거버넌스: 데이터 현지화·개인정보보호·국가간 전송 요건. [N]",
      "지식재산권: IP 라이선싱 의무, 표준필수특허(SEP) 고려사항. [N]",
      "수출 통제: 이중 용도 분류, 바세나르 체제 항목, 수출 제한. [N]"
    ]
  },
  "sec7": {
    "hype_cycle_phase": "혁신 촉발 / 과장 기대 최고점 / 환멸의 계곡 / 계몽의 경사로 / 생산성의 고원",
    "hype_cycle_plateau": "X년",
    "hype_cycle_source": "Gartner [연도] Hype Cycle for [카테고리]",
    "predictions": [
      {"year": "Year+1", "prediction": "예측 [N]", "confidence": "높음/중간/낮음", "implication": "시사점", "source": "IDC"},
      {"year": "Year+2", "prediction": "예측",     "confidence": "높음/중간/낮음", "implication": "시사점", "source": "IDC"},
      {"year": "Year+3", "prediction": "예측",     "confidence": "높음/중간/낮음", "implication": "시사점", "source": "IDC/Gartner"},
      {"year": "Year+4", "prediction": "예측",     "confidence": "높음/중간/낮음", "implication": "시사점", "source": "Gartner"},
      {"year": "Year+5", "prediction": "예측",     "confidence": "높음/중간/낮음", "implication": "시사점", "source": "McKinsey MGI"}
    ],
    "roadmap": [
      {"phase": "단기", "timeframe": "0~2년",  "milestone": "마일스톤 [N]", "trl": "TRL X", "enabler": "핵심 기술/정책", "risk": "주요 리스크"},
      {"phase": "중기", "timeframe": "2~5년",  "milestone": "마일스톤",     "trl": "TRL X", "enabler": "핵심 기술/정책", "risk": "주요 리스크"},
      {"phase": "장기", "timeframe": "5~10년", "milestone": "마일스톤",     "trl": "TRL X", "enabler": "핵심 기술/정책", "risk": "주요 리스크"}
    ]
  },
  "sec8": {
    "opportunities": [
      {"opportunity": "공동 연구개발",  "type": "연구",   "partners": "ETRI/KAIST/POSTECH", "funding": "IITP/BMBF",      "priority": "높음/중간/낮음", "timeline": "연도"},
      {"opportunity": "기술 이전",     "type": "상용화", "partners": "한국 중소기업/대기업", "funding": "MOTIE/KISTEP",  "priority": "높음/중간/낮음", "timeline": "연도"},
      {"opportunity": "정책 기여",     "type": "자문",   "partners": "MSIT/MOTIE",          "funding": "–",             "priority": "높음/중간/낮음", "timeline": "연도"},
      {"opportunity": "EU-한국 브리지","type": "협력",   "partners": "EU 연구 파트너",       "funding": "Horizon Europe","priority": "높음/중간/낮음", "timeline": "연도"}
    ],
    "risks": [
      {"risk": "리스크 1 [N]", "likelihood": "상/중/하", "impact": "상/중/하", "mitigation": "완화 방안", "owner": "담당자"},
      {"risk": "리스크 2",     "likelihood": "상/중/하", "impact": "상/중/하", "mitigation": "완화 방안", "owner": "담당자"},
      {"risk": "리스크 3",     "likelihood": "상/중/하", "impact": "상/중/하", "mitigation": "완화 방안", "owner": "담당자"}
    ],
    "actions": [
      "조치 1: ETRI/KISTEP과 구조화된 기술 브리핑 개최. [N]",
      "조치 2: Espacenet 및 KIPO를 통한 특허 자유 실시(FTO) 분석 착수.",
      "조치 3: 한국 연구 대학 또는 대기업 R&D 연구소 포함 Horizon Europe 파트너십 제안서 초안 작성."
    ]
  },
  "sec9": {
    "methodology": "자동화 데이터 수집 일정: Gartner·IDC·특허 DB 일일 API 수집; 키워드 목록·불리언 검색 문자열; 중복 제거·관련성 점수 로직; 월간 보고서 생성·QA 프로세스.",
    "quality": "출처 유형별 데이터 격차 식별; 주요 지표별 신뢰도(높음/중간/낮음); 업데이트 주기별 표시."
  }
}"""

_PROMPT_KO = """당신은 프라운호퍼 연구소 한국 사무소의 기술 시장 조사 분석가이다.
아래 {year}년 {month:02d}월 모니터링 기사 {count}건을 분석하여 구조화된 JSON 리포트를 작성하라.

인용 규칙 (엄격히 준수):
• 모든 사실적 주장, 통계, 트렌드 언급은 반드시 문장 끝에 [N] (N = 기사 ref 번호)을 붙인다.
• 제공된 데이터에 없는 내용을 만들거나 추론하지 말라.
• 근거를 제시할 수 없는 항목은 반드시 "N/A"로 기재한다.
• 수치·기업명·날짜는 원문과 100% 일치시킨다.
• 모든 서술은 '-함', '-임', '-전망됨', '-분석됨' 등 명사형으로 종결한다. '-습니다/ㅂ니다' 절대 사용 금지.
• 전문 약어(CAGR, TRL, LLM, GPU 등) 첫 등장 시 반드시 괄호 병기: 'CAGR(연평균 성장률)'.

대시보드 지표(metrics) 작성 규칙 (절대 위반 금지):
• 각 지표의 "value", "yoy", "forecast" 값은 제공된 기사에 수치가 명시된 경우에만 기입한다.
  명시되지 않은 경우 value는 "N/A", yoy는 "–", forecast는 "N/A"로 기재한다.
• "source" 필드에는 반드시 해당 수치를 직접 언급한 기사의 출처명(예: IEA, KEPCO, Financial Times Tech)을 그대로 사용한다.
  "Gartner/IDC", "KIET/KOTRA" 등 기사에 없는 출처를 임의로 넣지 말라.
• TRL, 특허 출원국, 시장점유율은 기사에 수치가 없으면 모두 N/A로 기재한다.
• 한 지표에 여러 기사 출처가 있으면 쉼표로 구분해 기입한다 (예: "IEA, KEPCO").

monthly_headline / monthly_context 작성 규칙:
• monthly_headline: 이달 기사 전체에서 가장 중요한 단일 사건이나 흐름을 1~2문장으로 서술. 구체적 수치 포함.
• monthly_context: 기사들이 공통으로 가리키는 구조적 변화를 3~4문장으로 서술. 독자가 이 단락만 읽어도 이달 핵심을 파악할 수 있어야 함.

sec1.key_findings 작성 규칙 (절대 준수):
• 반드시 정확히 4개 항목으로 작성하라: 시장 신호, 경쟁 신호, 한국 특화 신호, 리스크 신호.
• 각 항목은 기사 전체에서 해당 신호 유형에 가장 중요한 내용을 빠짐없이 담아야 한다.
• 시장 신호: 글로벌 시장 규모·투자·수요 관련 핵심 사실 (에너지·반도체·기업 등 모든 분야 포함).
• 경쟁 신호: 특정 기업이나 플레이어의 전략적 움직임 (시장 점유 확대, 설비 투자 가속 등).
• 한국 특화 신호: 한국 시장·정책·기업에 직접 관련된 내용.
• 리스크 신호: 공급망 위기, 규제 장벽, 기술적 병목 등 주요 리스크.
• 각 항목을 한 문장 안에 핵심 사실을 모두 담아 구체적으로 서술하라 (수치·기업명 포함). 내용을 빠뜨리거나 축약하지 말라.

sec2.definition 작성 규칙:
• 첫 문장에 반드시 technology_name 필드에 명시된 기술명을 그대로 사용하여 해당 기술이 무엇인지 명확히 정의하라.
• "이 기술은..." 또는 "AI와 에너지 관리 기술은..." 같이 모호하게 쓰지 말고, 기술명을 직접 명시하라.
• 예: "AI 기반 배터리 에너지 저장장치(BESS) 관리 기술은 ..."

N/A 처리 규칙:
• 근거 없이 작성할 수 없는 텍스트 필드(overview, definition 등)는 문장 형식으로 N/A를 언급하지 말고, 해당 필드 전체를 "N/A"로만 기재하라.
  잘못된 예: "글로벌 시장 규모는 N/A이다 [N/A]."
  올바른 예: "N/A"
• trl_table의 assessment에 실제 TRL 단계(TRL 1~9)를 알 수 없으면 반드시 "N/A"로 기재하라. "TRL X"나 "X년" 같은 템플릿 문자는 절대 그대로 남기지 말라.

출력: 마크다운 없이 유효한 JSON만 반환할 것. 아래 스키마를 정확히 따를 것:
{schema}

인덱스된 기사 데이터 (JSON):
{articles_json}"""


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_element = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run_element.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


_NA_VALUES = {"n/a", "–", "-", "", "na", "n.a.", "not available", "unknown"}

# Template placeholder patterns that the LLM sometimes leaves unfilled
_PLACEHOLDER_RE = re.compile(
    r"^(?:TRL\s*[Xx\?0]|[Xx]\s*년|[Xx]\s*years?|[Xx]\s*yr|기술명|tech(?:nology)?|"
    r"driver\s*\d+|barrier\s*\d+|[Xx]\s*%|[Xx]\s*(?:billion|million|B|M)|"
    r"\$\s*X[BM]?|\€\s*X[BM]?|¥\s*X[BM]?|₩\s*X[BM]?)$",
    re.IGNORECASE,
)


def _is_na(val: Any) -> bool:
    """Return True if a value is effectively empty, N/A, or an unfilled template placeholder."""
    s = str(val).strip()
    return s.lower() in _NA_VALUES or bool(_PLACEHOLDER_RE.match(s))


def _has_real_data(items: list[dict], keys: list[str]) -> bool:
    """Return True if at least one item has a meaningful value for any of the given keys."""
    for item in items:
        for key in keys:
            if not _is_na(item.get(key, "")):
                return True
    return False


def _has_real_text(text: str) -> bool:
    """Return True if a text field is not empty / N/A."""
    return not _is_na(text)


def _has_real_list(items: list[str]) -> bool:
    """Return True if a list contains at least one non-N/A entry."""
    return any(not _is_na(item) for item in items)


def _bold_cell(cell, text: str) -> None:
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    style: str = "Light Grid Accent 1",
) -> None:
    tbl = document.add_table(rows=1, cols=len(headers))
    tbl.style = style
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _bold_cell(hdr[i], h)
    for row_vals in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_vals):
            if i < len(cells):
                cells[i].text = str(val)
    document.add_paragraph()


def _extract_json(raw: str) -> dict:
    """Strip optional markdown fences and parse JSON."""
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


# ---------------------------------------------------------------------------
# ReportGenerator
# ---------------------------------------------------------------------------


class ReportGenerator:
    def __init__(self, settings: Settings) -> None:
        load_dotenv(PROJECT_ROOT / ".env")

        if not os.getenv("OPENAI_API_KEY"):
            raise ValueError("OPENAI_API_KEY is required for monthly reporting")

        self.llm = OpenAI(
            api_key=os.getenv("OPENAI_API_KEY"),
            base_url=os.getenv("OPENAI_BASE_URL"),
        )
        self.output_dir = settings.reports_output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # LLM synthesis
    # ------------------------------------------------------------------

    def _synthesize(self, year: int, month: int, logs: list[dict], lang: str = "en") -> dict:
        """Call LLM and return the structured report dict."""
        indexed_logs = [
            {
                "ref": i,
                "date": item["log_date"],
                "title": item["title"],
                "url": item["url"],
                "source": item.get("source_name", item["category"]),
                "category": item["category"],
                "summary": item.get("ko_summary", item["llm_summary"]) if lang == "ko" else item["llm_summary"],
                "trends": item["key_trends"],
            }
            for i, item in enumerate(logs, start=1)
        ]

        if lang == "ko":
            prompt = _PROMPT_KO.format(
                year=year,
                month=month,
                count=len(logs),
                schema=_SCHEMA_JSON_KO,
                articles_json=json.dumps(indexed_logs, ensure_ascii=False),
            )
            system_msg = (
                "당신은 간결하고 인용이 풍부한 한국어 기술 시장 조사 리포트를 작성한다. "
                "모든 사실적 문장 끝에 [N] 인텍스트 인용을 반드시 붙인다. "
                "소스 데이터에 없는 정보는 절대 추가하지 않는다. "
                "'-습니다/ㅂ니다' 어미 사용 절대 금지. 명사형 종결 필수. "
                "전문 약어 첫 등장 시 괄호 병기 필수."
            )
        else:
            prompt = _PROMPT_EN.format(
                year=year,
                month=month,
                count=len(logs),
                schema=_SCHEMA_JSON,
                articles_json=json.dumps(indexed_logs, ensure_ascii=False),
            )
            system_msg = (
                "You write concise, citation-heavy technology market research reports. "
                "Attach [N] in-text citations at every factual sentence. "
                "Never fabricate information not present in the source data. "
                "Return only valid JSON – no markdown, no explanation."
            )

        response = self.llm.chat.completions.create(
            model=os.getenv("MODEL_NAME", "gemini-1.5-flash"),
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
        )
        raw = (response.choices[0].message.content or "").strip()
        try:
            return _extract_json(raw)
        except Exception as exc:
            logger.error("JSON parse failed: %s\nRaw output (first 500 chars): %s", exc, raw[:500])
            raise

    # ------------------------------------------------------------------
    # Document builders
    # ------------------------------------------------------------------

    @staticmethod
    def _setup_document() -> Document:
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        return document

    @staticmethod
    def _add_cover(document: Document, title: str, subtitle: str, fields: list[tuple[str, str]]) -> None:
        """Add a cover page with title, subtitle, and a meta-data table."""
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = _HEADER_BLUE

        sub = document.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(subtitle)
        sr.font.size = Pt(14)
        sr.font.color.rgb = _SUBTITLE_BLUE

        document.add_paragraph()

        meta_tbl = document.add_table(rows=len(fields), cols=2)
        meta_tbl.style = "Light Grid Accent 1"
        for i, (label, value) in enumerate(fields):
            row = meta_tbl.rows[i].cells
            _bold_cell(row[0], label)
            row[1].text = value

        document.add_page_break()

    def _build_document(
        self, year: int, month: int, logs: list[dict], d: dict
    ) -> Document:
        """Build English TMR document matching the Fraunhofer template structure."""
        document = self._setup_document()

        # ── COVER ──────────────────────────────────────────────────────────
        self._add_cover(
            document,
            "Technology Market Research Report",
            "Fraunhofer Institute | Korea Office",
            [
                ("Technology Name", d.get("technology_name", "Tech Market Intelligence")),
                ("Report Period", f"{year}-{month:02d}"),
                ("Prepared by", "Fraunhofer Korea – Automated Monitoring System"),
                ("Version", "v1.0"),
                ("Classification", "Internal / Confidential"),
            ],
        )

        s1 = d.get("sec1", {})
        s2 = d.get("sec2", {})
        s3 = d.get("sec3", {})
        s4 = d.get("sec4", {})
        s5 = d.get("sec5", {})
        s6 = d.get("sec6", {})
        s7 = d.get("sec7", {})
        s8 = d.get("sec8", {})
        s9 = d.get("sec9", {})
        tech_name_en = d.get("technology_name", "")

        # ── MONTHLY KEY TAKEAWAYS ──────────────────────────────────────────
        document.add_heading("This Month's Key Takeaways", level=1)

        headline_en = d.get("monthly_headline", "")
        if headline_en:
            hl_para = document.add_paragraph()
            hl_run = hl_para.add_run(headline_en)
            hl_run.bold = True
            hl_run.font.size = Pt(12)
            hl_run.font.color.rgb = _HEADER_BLUE

        monthly_ctx_en = d.get("monthly_context", "")
        if monthly_ctx_en:
            document.add_paragraph(monthly_ctx_en)

        document.add_paragraph()
        document.add_paragraph("▶ Key Findings").bold = True
        for finding in s1.get("key_findings", []):
            document.add_paragraph(finding, style="List Bullet")

        document.add_page_break()

        # ── Section counter (incremented each time a section is rendered) ──
        _sn = [0]

        # ── SECTION 1: EXECUTIVE BRIEF ─────────────────────────────────────
        _sn[0] += 1
        sn = _sn[0]
        document.add_heading(f"SECTION {sn} – EXECUTIVE BRIEF", level=1)
        _ssn = [0]

        if _has_real_text(s1.get("snapshot", "")):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} Technology Snapshot", level=2)
            document.add_paragraph(s1.get("snapshot"))

        en_findings = [f for f in s1.get("key_findings", []) if not _is_na(f)]
        if en_findings:
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} Key Findings", level=2)
            for finding in en_findings:
                document.add_paragraph(finding, style="List Bullet")

        metrics_en = s1.get("metrics", [])
        if _has_real_data(metrics_en, ["value"]):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} Metrics Dashboard", level=2)
            _add_table(
                document,
                ["Metric", "Current Value", "YoY Change", "Forecast (3yr)", "Source"],
                [
                    [r.get("metric",""), r.get("value",""), r.get("yoy",""), r.get("forecast",""), r.get("source","")]
                    for r in metrics_en
                ],
            )

        # ── SECTION 2: TECHNOLOGY PROFILE ─────────────────────
        definition_en = s2.get("definition", "")
        trl_rows_en = s2.get("trl_table", [])
        comparison_en = s2.get("comparison_table", [])
        patents_en = s2.get("patents", [])

        has_def_en = _has_real_text(definition_en)
        has_trl_en = _has_real_data(trl_rows_en, ["assessment"])
        has_cmp_en = _has_real_data(comparison_en, ["this_tech"])
        has_pat_en = _has_real_data(patents_en, ["no", "title"])

        if any([has_def_en, has_trl_en, has_cmp_en, has_pat_en]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"SECTION {sn} – TECHNOLOGY PROFILE", level=1)
            _ssn = [0]

            if has_def_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Technology Definition & Principles", level=2)
                document.add_paragraph(definition_en)
                document.add_paragraph("Sources: IEEE Xplore, Fraunhofer-Publica, ETRI, arXiv").italic = True

            if has_trl_en:
                _ssn[0] += 1
                trl_heading_en = f"{sn}.{_ssn[0]} Technology Readiness Level (TRL) – {tech_name_en}" if tech_name_en else f"{sn}.{_ssn[0]} Technology Readiness Level (TRL)"
                document.add_heading(trl_heading_en, level=2)
                _add_table(
                    document,
                    ["Dimension", "Assessment", "Basis"],
                    [[r.get("dimension",""), r.get("assessment",""), r.get("basis","")] for r in trl_rows_en],
                )

            if has_cmp_en or _has_real_text(s2.get("differentiation", "")):
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Technology Differentiation", level=2)
                if _has_real_text(s2.get("differentiation", "")):
                    document.add_paragraph(s2.get("differentiation"))
                if has_cmp_en:
                    alt_a = s2.get("alt_a_name", "Alternative A")
                    alt_b = s2.get("alt_b_name", "Alternative B")
                    _add_table(
                        document,
                        ["Feature", "This Technology", alt_a, alt_b],
                        [
                            [r.get("feature",""), r.get("this_tech",""), r.get("alt_a",""), r.get("alt_b","")]
                            for r in comparison_en
                        ],
                    )

            if has_pat_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Patent Landscape", level=2)
                document.add_paragraph("Sources: Espacenet (EPO), KIPO, DPMA, JPO, CNIPA, Google Patents").italic = True
                _add_table(
                    document,
                    ["Patent No.", "Title", "Assignee", "Country", "Year", "Source"],
                    [
                        [p.get("no",""), p.get("title",""), p.get("assignee",""), p.get("country",""), p.get("year",""), p.get("source","")]
                        for p in patents_en
                    ],
                )

        # ── SECTION 3: MARKET ANALYSIS ─────────────────────────
        overview_en = s3.get("overview", "")
        seg_en = s3.get("segmentation", [])
        reg_en = s3.get("regional", [])
        db_en = s3.get("drivers_barriers", [])

        has_ov_en = _has_real_text(overview_en)
        has_seg_en = _has_real_data(seg_en, ["size", "name"])
        has_reg_en = _has_real_data(reg_en, ["size", "drivers"])
        has_db_en = _has_real_data(db_en, ["driver", "barrier"])

        if any([has_ov_en, has_seg_en, has_reg_en, has_db_en]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"SECTION {sn} – MARKET ANALYSIS", level=1)
            _ssn = [0]

            if has_ov_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Market Overview", level=2)
                document.add_paragraph(overview_en)

            if has_seg_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Market Segmentation", level=2)
                _add_table(
                    document,
                    ["Segment Type", "Segment Name", "Size", "Share", "Growth", "Notes"],
                    [
                        [r.get("type",""), r.get("name",""), r.get("size",""), r.get("share",""), r.get("growth",""), r.get("notes","")]
                        for r in seg_en
                    ],
                )

            if has_reg_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Regional Deep-Dive", level=2)
                _add_table(
                    document,
                    ["Region", "Market Size", "Key Drivers", "Policy Environment", "Top Players", "Source"],
                    [
                        [r.get("region",""), r.get("size",""), r.get("drivers",""), r.get("policy",""), r.get("players",""), r.get("source","")]
                        for r in reg_en
                    ],
                )

            if has_db_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Market Drivers & Barriers", level=2)
                _add_table(
                    document,
                    ["Drivers (with Source)", "Barriers (with Source)"],
                    [[r.get("driver",""), r.get("barrier","")] for r in db_en],
                )

        # ── SECTION 4: COMPETITIVE LANDSCAPE ──────────────────────
        vendors_en = s4.get("vendors", [])
        korea_ctx_en = s4.get("korea_context", "")
        swot_en = s4.get("swot", {})
        ff_en = s4.get("five_forces", [])

        sw_en = swot_en.get("strengths",[]) + swot_en.get("weaknesses",[]) + swot_en.get("opportunities",[]) + swot_en.get("threats",[])
        has_ven_en = _has_real_data(vendors_en, ["vendor"])
        has_kctx_en = _has_real_text(korea_ctx_en)
        has_swot_en = _has_real_list(sw_en)
        has_ff_en = _has_real_data(ff_en, ["intensity"])

        if any([has_ven_en, has_kctx_en, has_swot_en, has_ff_en]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"SECTION {sn} – COMPETITIVE LANDSCAPE", level=1)
            _ssn = [0]

            if has_ven_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Vendor Mapping", level=2)
                _add_table(
                    document,
                    ["Vendor", "HQ", "Type", "Core Offering", "Strategy", "Market Position", "Source"],
                    [
                        [v.get("vendor",""), v.get("hq",""), v.get("type",""), v.get("offering",""), v.get("strategy",""), v.get("position",""), v.get("source","")]
                        for v in vendors_en
                    ],
                )

            if has_kctx_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Korea-Specific Competitive Context", level=2)
                document.add_paragraph(korea_ctx_en)

            if has_swot_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} SWOT Analysis", level=2)
                strengths_en = "\n".join(s for s in swot_en.get("strengths",[]) if not _is_na(s))
                weaknesses_en = "\n".join(s for s in swot_en.get("weaknesses",[]) if not _is_na(s))
                opps_en = "\n".join(s for s in swot_en.get("opportunities",[]) if not _is_na(s))
                threats_en = "\n".join(s for s in swot_en.get("threats",[]) if not _is_na(s))
                if strengths_en or weaknesses_en:
                    _add_table(document, ["Strengths", "Weaknesses"], [[strengths_en, weaknesses_en]])
                if opps_en or threats_en:
                    _add_table(document, ["Opportunities", "Threats"], [[opps_en, threats_en]])

            if has_ff_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Porter's Five Forces Summary", level=2)
                _add_table(
                    document,
                    ["Force", "Intensity", "Key Factor", "Implication"],
                    [
                        [f.get("force",""), f.get("intensity",""), f.get("key_factor",""), f.get("implication","")]
                        for f in ff_en
                    ],
                )

        # ── SECTION 5: INNOVATION & R&D ────────────────────────────
        pubs_en = s5.get("publications", [])
        fund_en = s5.get("funding", [])
        dirs_en = [d for d in s5.get("emerging_directions", []) if not _is_na(d)]

        has_pubs_en = _has_real_data(pubs_en, ["title"])
        has_fund_en = _has_real_data(fund_en, ["program", "budget"])
        has_dirs_en = bool(dirs_en)

        if any([has_pubs_en, has_fund_en, has_dirs_en]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"SECTION {sn} – INNOVATION & R&D LANDSCAPE", level=1)
            _ssn = [0]

            if has_pubs_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Academic Publication Trends", level=2)
                document.add_paragraph("Sources: IEEE Xplore, ACM Digital Library, arXiv, Springer/Elsevier, ETRI Journal, KISTI").italic = True
                _add_table(
                    document,
                    ["Title", "Authors", "Journal / Conference", "Year", "Citations", "DOI"],
                    [
                        [p.get("title",""), p.get("authors",""), p.get("journal",""), p.get("year",""), p.get("citations",""), p.get("doi","")]
                        for p in pubs_en
                    ],
                )

            if has_fund_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} R&D Funding Programs", level=2)
                _add_table(
                    document,
                    ["Program", "Funding Body", "Region", "Budget", "Focus Area", "Timeline", "Source"],
                    [
                        [r.get("program",""), r.get("body",""), r.get("region",""), r.get("budget",""), r.get("focus",""), r.get("timeline",""), r.get("source","")]
                        for r in fund_en
                    ],
                )

            if has_dirs_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Emerging Research Directions", level=2)
                for direction in dirs_en:
                    document.add_paragraph(direction, style="List Bullet")

        # ── SECTION 6: REGULATORY & POLICY ────────────────────────
        pol_en = s6.get("policies", [])
        comp_en = [c for c in s6.get("compliance", []) if not _is_na(c)]

        has_pol_en = _has_real_data(pol_en, ["policy"])
        has_comp_en = bool(comp_en)

        if any([has_pol_en, has_comp_en]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"SECTION {sn} – REGULATORY & POLICY ENVIRONMENT", level=1)
            _ssn = [0]

            if has_pol_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Policy Tracker", level=2)
                _add_table(
                    document,
                    ["Policy / Regulation", "Issuing Body", "Region", "Status", "Effective Date", "Impact", "Source"],
                    [
                        [p.get("policy",""), p.get("body",""), p.get("region",""), p.get("status",""), p.get("date",""), p.get("impact",""), p.get("source","")]
                        for p in pol_en
                    ],
                )

            if has_comp_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Compliance Considerations", level=2)
                for item in comp_en:
                    document.add_paragraph(item, style="List Bullet")

        # ── SECTION 7: TECHNOLOGY FORECAST ────────────────────────
        hc_phase_en = s7.get("hype_cycle_phase", "")
        pred_en = s7.get("predictions", [])
        road_en = s7.get("roadmap", [])

        has_hc_en = _has_real_text(hc_phase_en)
        has_pred_en = _has_real_data(pred_en, ["prediction"])
        has_road_en = _has_real_data(road_en, ["milestone"])

        if any([has_hc_en, has_pred_en, has_road_en]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"SECTION {sn} – TECHNOLOGY FORECAST", level=1)
            _ssn = [0]

            if has_hc_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Gartner Hype Cycle Position", level=2)
                hc = document.add_table(rows=3, cols=2)
                hc.style = "Light Grid Accent 1"
                _bold_cell(hc.rows[0].cells[0], "Current Phase")
                hc.rows[0].cells[1].text = hc_phase_en
                _bold_cell(hc.rows[1].cells[0], "Time to Plateau")
                hc.rows[1].cells[1].text = s7.get("hype_cycle_plateau", "N/A")
                _bold_cell(hc.rows[2].cells[0], "Source")
                hc.rows[2].cells[1].text = s7.get("hype_cycle_source", "N/A")
                document.add_paragraph()

            if has_pred_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} IDC-Style 5-Year Predictions", level=2)
                _add_table(
                    document,
                    ["Year", "Prediction", "Confidence", "Implication", "Source"],
                    [
                        [r.get("year",""), r.get("prediction",""), r.get("confidence",""), r.get("implication",""), r.get("source","")]
                        for r in pred_en
                    ],
                )

            if has_road_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Technology Roadmap", level=2)
                _add_table(
                    document,
                    ["Phase", "Timeframe", "Milestone", "TRL", "Key Enabler", "Risk"],
                    [
                        [r.get("phase",""), r.get("timeframe",""), r.get("milestone",""), r.get("trl",""), r.get("enabler",""), r.get("risk","")]
                        for r in road_en
                    ],
                )

        # ── SECTION 8: STRATEGIC IMPLICATIONS ───────────────────────
        opps8_en = s8.get("opportunities", [])
        risks8_en = s8.get("risks", [])
        act8_en = [a for a in s8.get("actions", []) if not _is_na(a)]

        has_opps8_en = _has_real_data(opps8_en, ["opportunity"])
        has_risks8_en = _has_real_data(risks8_en, ["risk"])
        has_act8_en = bool(act8_en)

        if any([has_opps8_en, has_risks8_en, has_act8_en]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"SECTION {sn} – STRATEGIC IMPLICATIONS FOR FRAUNHOFER KOREA", level=1)
            _ssn = [0]

            if has_opps8_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Opportunity Assessment", level=2)
                _add_table(
                    document,
                    ["Opportunity", "Type", "Potential Partners", "Funding Source", "Priority", "Timeline"],
                    [
                        [o.get("opportunity",""), o.get("type",""), o.get("partners",""), o.get("funding",""), o.get("priority",""), o.get("timeline","")]
                        for o in opps8_en
                    ],
                )

            if has_risks8_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Risk Register", level=2)
                _add_table(
                    document,
                    ["Risk", "Likelihood", "Impact", "Mitigation", "Owner"],
                    [
                        [r.get("risk",""), r.get("likelihood",""), r.get("impact",""), r.get("mitigation",""), r.get("owner","")]
                        for r in risks8_en
                    ],
                )

            if has_act8_en:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} Recommended Actions (Next 90 Days)", level=2)
                for action in act8_en:
                    document.add_paragraph(action, style="List Bullet")

        # ── SECTION 9: SOURCES & METHODOLOGY ───────────────────────
        _sn[0] += 1
        sn = _sn[0]
        document.add_heading(f"SECTION {sn} – SOURCES & METHODOLOGY", level=1)
        _ssn = [0]

        _ssn[0] += 1
        document.add_heading(f"{sn}.{_ssn[0]} Source Registry", level=2)
        document.add_paragraph("Academic & Research:").bold = True
        for src in ["IEEE Xplore – ieee.org/xplore", "ACM Digital Library – dl.acm.org",
                    "arXiv – arxiv.org", "Fraunhofer-Publica – publica.fraunhofer.de",
                    "ETRI Journal – etri.re.kr", "KISTI – kisti.re.kr",
                    "Springer / Elsevier / Wiley – via institutional access"]:
            document.add_paragraph(src, style="List Bullet")

        document.add_paragraph("Market Intelligence:").bold = True
        for src in ["Gartner – gartner.com", "IDC – idc.com", "Statista – statista.com",
                    "McKinsey Global Institute – mckinsey.com/mgi", "KIET – kiet.re.kr",
                    "KOTRA – kotra.or.kr", "KISDI – kisdi.re.kr", "ADB – adb.org",
                    "OECD iLibrary – oecd-ilibrary.org", "Eurostat – ec.europa.eu/eurostat",
                    "Destatis – destatis.de", "KOSTAT – kostat.go.kr"]:
            document.add_paragraph(src, style="List Bullet")

        document.add_paragraph("Patent & IP:").bold = True
        for src in ["Espacenet (EPO) – espacenet.epo.org", "KIPO – kipo.go.kr",
                    "DPMA – dpma.de", "JPO – j-platpat.inpit.go.jp",
                    "CNIPA – cnipa.gov.cn", "Google Patents – patents.google.com"]:
            document.add_paragraph(src, style="List Bullet")

        document.add_paragraph("Government & Policy:").bold = True
        for src in ["EU Commission – ec.europa.eu", "BMBF – bmbf.de",
                    "MOTIE – motie.go.kr", "MSIT – msit.go.kr",
                    "KISTEP – kistep.re.kr", "IITP – iitp.kr",
                    "NIST – nist.gov", "IEA – iea.org", "NEDO – nedo.go.jp",
                    "METI – meti.go.jp", "APEC – apec.org",
                    "NITI Aayog – niti.gov.in", "EDB Singapore – edb.gov.sg"]:
            document.add_paragraph(src, style="List Bullet")

        meth_en = s9.get("methodology", "")
        if _has_real_text(meth_en):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} Automated Monitoring Methodology", level=2)
            document.add_paragraph(meth_en)

        qual_en = s9.get("quality", "")
        if _has_real_text(qual_en):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} Data Quality & Limitations", level=2)
            document.add_paragraph(qual_en)

        # ── REFERENCES (from actual article logs) ──────────────────────────
        document.add_page_break()
        document.add_heading("References", level=1)
        accessed = date.today().isoformat()
        _add_table(
            document,
            ["No.", "Source", "Title", "Date", "URL"],
            [
                [
                    f"[{i}]",
                    item.get("source_name", item.get("category", "Unknown")),
                    item["title"],
                    item["log_date"],
                    item["url"],
                ]
                for i, item in enumerate(logs, start=1)
            ],
        )

        # ── SOURCE INDEX ───────────────────────────────────────────────────
        document.add_heading("Source Index", level=1)
        src_tbl = document.add_table(rows=1, cols=4)
        src_tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["Date", "Title", "Category", "Source URL"]):
            _bold_cell(src_tbl.rows[0].cells[i], h)
        for item in logs:
            row = src_tbl.add_row().cells
            row[0].text = item["log_date"]
            row[1].text = item["title"]
            row[2].text = CATEGORY_LABELS.get(item["category"], item["category"])
            _add_hyperlink(row[3].paragraphs[0], item["url"], item["url"])

        document.add_paragraph(f"\n— End of Report — Generated {accessed}")
        return document

    def _build_document_ko(
        self, year: int, month: int, logs: list[dict], d: dict
    ) -> Document:
        """Build Korean TMR document matching the Fraunhofer template structure."""
        document = self._setup_document()

        # ── 표지 ───────────────────────────────────────────────────────────
        self._add_cover(
            document,
            "기술 시장 조사 보고서",
            "프라운호퍼 연구소 | 한국 사무소",
            [
                ("기술명", d.get("technology_name", "기술 시장 인텔리전스")),
                ("보고서 기간", f"{year}년 {month:02d}월"),
                ("작성자", "프라운호퍼 한국 – 자동 모니터링 시스템"),
                ("버전", "v1.0"),
                ("분류", "내부용 / 기밀"),
            ],
        )

        s1 = d.get("sec1", {})
        s2 = d.get("sec2", {})
        s3 = d.get("sec3", {})
        s4 = d.get("sec4", {})
        s5 = d.get("sec5", {})
        s6 = d.get("sec6", {})
        s7 = d.get("sec7", {})
        s8 = d.get("sec8", {})
        s9 = d.get("sec9", {})
        tech_name = d.get("technology_name", "")

        # ── 이달의 핵심 내용 (Executive Summary) ──────────────────────────
        document.add_heading("이달의 핵심 내용", level=1)

        headline = d.get("monthly_headline", "")
        if headline:
            hl_para = document.add_paragraph()
            hl_run = hl_para.add_run(headline)
            hl_run.bold = True
            hl_run.font.size = Pt(12)
            hl_run.font.color.rgb = _HEADER_BLUE

        monthly_ctx = d.get("monthly_context", "")
        if monthly_ctx:
            document.add_paragraph(monthly_ctx)

        document.add_paragraph()
        document.add_paragraph("▶ 주요 발견 사항").bold = True
        for finding in s1.get("key_findings", []):
            document.add_paragraph(finding, style="List Bullet")

        document.add_page_break()

        # ── 섹션 카운터 (섹션이 렌더링될 때마다 증가) ────────────────────
        _sn = [0]

        # ── 섹션 1: 핵심 요약 ──────────────────────────────────────────────
        _sn[0] += 1
        sn = _sn[0]
        document.add_heading(f"섹션 {sn} – 핵심 요약", level=1)
        _ssn = [0]

        if _has_real_text(s1.get("snapshot", "")):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} 기술 개요", level=2)
            document.add_paragraph(s1.get("snapshot"))

        findings = [f for f in s1.get("key_findings", []) if not _is_na(f)]
        if findings:
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} 주요 발견 사항", level=2)
            for finding in findings:
                document.add_paragraph(finding, style="List Bullet")

        metrics = s1.get("metrics", [])
        if _has_real_data(metrics, ["value"]):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} 핵심 지표 대시보드", level=2)
            _add_table(
                document,
                ["지표", "현재 값", "전년 대비 변동", "예측 (3년)", "출처"],
                [
                    [r.get("metric",""), r.get("value",""), r.get("yoy",""), r.get("forecast",""), r.get("source","")]
                    for r in metrics
                ],
            )

        # ── 섹션 2: 기술 프로파일 ─────────────────────────────────────────
        definition = s2.get("definition", "")
        comparison = s2.get("comparison_table", [])
        trl_rows = s2.get("trl_table", [])
        patents = s2.get("patents", [])

        has_definition = _has_real_text(definition)
        has_trl = _has_real_data(trl_rows, ["assessment"])
        has_comparison = _has_real_data(comparison, ["this_tech"])
        has_patents = _has_real_data(patents, ["no", "title"])

        if any([has_definition, has_trl, has_comparison, has_patents]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"섹션 {sn} – 기술 프로파일", level=1)
            _ssn = [0]

            if has_definition:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 기술 정의 및 원리", level=2)
                document.add_paragraph(definition)
                document.add_paragraph("출처: IEEE Xplore, Fraunhofer-Publica, ETRI, arXiv").italic = True

            if has_trl:
                _ssn[0] += 1
                trl_heading = f"{sn}.{_ssn[0]} {tech_name} 기술 준비 단계 (TRL)" if tech_name else f"{sn}.{_ssn[0]} 기술 준비 단계 (TRL)"
                document.add_heading(trl_heading, level=2)
                _add_table(
                    document,
                    ["평가 항목", "평가 결과", "평가 근거"],
                    [[r.get("dimension",""), r.get("assessment",""), r.get("basis","")] for r in trl_rows],
                )

            if has_comparison or _has_real_text(s2.get("differentiation", "")):
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 기술 차별화", level=2)
                if _has_real_text(s2.get("differentiation", "")):
                    document.add_paragraph(s2.get("differentiation"))
                if has_comparison:
                    alt_a = s2.get("alt_a_name", "대안 기술 A")
                    alt_b = s2.get("alt_b_name", "대안 기술 B")
                    _add_table(
                        document,
                        ["항목", "해당 기술", alt_a, alt_b],
                        [
                            [r.get("feature",""), r.get("this_tech",""), r.get("alt_a",""), r.get("alt_b","")]
                            for r in comparison
                        ],
                    )

            if has_patents:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 특허 현황", level=2)
                document.add_paragraph("출처: Espacenet (EPO), KIPO, DPMA, JPO, CNIPA, Google Patents").italic = True
                _add_table(
                    document,
                    ["특허 번호", "제목", "출원인", "국가", "연도", "출처"],
                    [
                        [p.get("no",""), p.get("title",""), p.get("assignee",""), p.get("country",""), p.get("year",""), p.get("source","")]
                        for p in patents
                    ],
                )

        # ── 섹션 3: 시장 분석 ─────────────────────────────────────────────
        overview = s3.get("overview", "")
        segmentation = s3.get("segmentation", [])
        regional = s3.get("regional", [])
        drivers_barriers = s3.get("drivers_barriers", [])

        has_overview = _has_real_text(overview)
        has_seg = _has_real_data(segmentation, ["size", "name"])
        has_regional = _has_real_data(regional, ["size", "drivers"])
        has_db = _has_real_data(drivers_barriers, ["driver", "barrier"])

        if any([has_overview, has_seg, has_regional, has_db]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"섹션 {sn} – 시장 분석", level=1)
            _ssn = [0]

            if has_overview:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 시장 개요", level=2)
                document.add_paragraph(overview)

            if has_seg:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 시장 세분화", level=2)
                _add_table(
                    document,
                    ["세분화 유형", "세그먼트명", "규모", "점유율", "성장률", "비고"],
                    [
                        [r.get("type",""), r.get("name",""), r.get("size",""), r.get("share",""), r.get("growth",""), r.get("notes","")]
                        for r in segmentation
                    ],
                )

            if has_regional:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 지역별 심층 분석", level=2)
                _add_table(
                    document,
                    ["지역", "시장 규모", "주요 동인", "정책 환경", "주요 플레이어", "출처"],
                    [
                        [r.get("region",""), r.get("size",""), r.get("drivers",""), r.get("policy",""), r.get("players",""), r.get("source","")]
                        for r in regional
                    ],
                )

            if has_db:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 시장 동인 및 장벽", level=2)
                _add_table(
                    document,
                    ["동인 (출처 포함)", "장벽 (출처 포함)"],
                    [[r.get("driver",""), r.get("barrier","")] for r in drivers_barriers],
                )

        # ── 섹션 4: 경쟁 환경 ─────────────────────────────────────────────
        vendors = s4.get("vendors", [])
        korea_ctx = s4.get("korea_context", "")
        swot = s4.get("swot", {})
        five_forces = s4.get("five_forces", [])

        sw_items = swot.get("strengths", []) + swot.get("weaknesses", []) + swot.get("opportunities", []) + swot.get("threats", [])
        has_vendors = _has_real_data(vendors, ["vendor"])
        has_korea_ctx = _has_real_text(korea_ctx)
        has_swot = _has_real_list(sw_items)
        has_five_forces = _has_real_data(five_forces, ["intensity"])

        if any([has_vendors, has_korea_ctx, has_swot, has_five_forces]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"섹션 {sn} – 경쟁 환경", level=1)
            _ssn = [0]

            if has_vendors:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 벤더 매핑", level=2)
                _add_table(
                    document,
                    ["벤더", "본사", "유형", "핵심 제공 서비스", "전략", "시장 포지션", "출처"],
                    [
                        [v.get("vendor",""), v.get("hq",""), v.get("type",""), v.get("offering",""), v.get("strategy",""), v.get("position",""), v.get("source","")]
                        for v in vendors
                    ],
                )

            if has_korea_ctx:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 한국 특화 경쟁 환경", level=2)
                document.add_paragraph(korea_ctx)

            if has_swot:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} SWOT 분석", level=2)
                strengths = "\n".join(s for s in swot.get("strengths", []) if not _is_na(s))
                weaknesses = "\n".join(s for s in swot.get("weaknesses", []) if not _is_na(s))
                opportunities = "\n".join(s for s in swot.get("opportunities", []) if not _is_na(s))
                threats = "\n".join(s for s in swot.get("threats", []) if not _is_na(s))
                if strengths or weaknesses:
                    _add_table(document, ["강점 (Strengths)", "약점 (Weaknesses)"], [[strengths, weaknesses]])
                if opportunities or threats:
                    _add_table(document, ["기회 (Opportunities)", "위협 (Threats)"], [[opportunities, threats]])

            if has_five_forces:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 포터의 5가지 경쟁 요인 요약", level=2)
                _add_table(
                    document,
                    ["경쟁 요인", "강도", "핵심 요인", "시사점"],
                    [
                        [f.get("force",""), f.get("intensity",""), f.get("key_factor",""), f.get("implication","")]
                        for f in five_forces
                    ],
                )

        # ── 섹션 5: 혁신 및 R&D ───────────────────────────────────────────
        publications = s5.get("publications", [])
        funding = s5.get("funding", [])
        directions = [d for d in s5.get("emerging_directions", []) if not _is_na(d)]

        has_pubs = _has_real_data(publications, ["title"])
        has_funding = _has_real_data(funding, ["program", "budget"])
        has_directions = bool(directions)

        if any([has_pubs, has_funding, has_directions]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"섹션 {sn} – 혁신 및 연구개발(R&D) 현황", level=1)
            _ssn = [0]

            if has_pubs:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 학술 논문 발표 동향", level=2)
                document.add_paragraph("출처: IEEE Xplore, ACM Digital Library, arXiv, Springer/Elsevier, ETRI Journal, KISTI").italic = True
                _add_table(
                    document,
                    ["제목", "저자", "저널 / 학회", "연도", "인용 수", "DOI"],
                    [
                        [p.get("title",""), p.get("authors",""), p.get("journal",""), p.get("year",""), p.get("citations",""), p.get("doi","")]
                        for p in publications
                    ],
                )

            if has_funding:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 연구개발 지원 프로그램", level=2)
                _add_table(
                    document,
                    ["프로그램", "지원 기관", "지역", "예산", "핵심 분야", "기간", "출처"],
                    [
                        [r.get("program",""), r.get("body",""), r.get("region",""), r.get("budget",""), r.get("focus",""), r.get("timeline",""), r.get("source","")]
                        for r in funding
                    ],
                )

            if has_directions:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 신흥 연구 방향", level=2)
                for direction in directions:
                    document.add_paragraph(direction, style="List Bullet")

        # ── 섹션 6: 규제 및 정책 환경 ────────────────────────────────────
        policies = s6.get("policies", [])
        compliance = [c for c in s6.get("compliance", []) if not _is_na(c)]

        has_policies = _has_real_data(policies, ["policy"])
        has_compliance = bool(compliance)

        if any([has_policies, has_compliance]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"섹션 {sn} – 규제 및 정책 환경", level=1)
            _ssn = [0]

            if has_policies:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 정책 트래커", level=2)
                _add_table(
                    document,
                    ["정책 / 규제", "발행 기관", "지역", "상태", "시행 일자", "기술 영향", "출처"],
                    [
                        [p.get("policy",""), p.get("body",""), p.get("region",""), p.get("status",""), p.get("date",""), p.get("impact",""), p.get("source","")]
                        for p in policies
                    ],
                )

            if has_compliance:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 컴플라이언스 고려 사항", level=2)
                for item in compliance:
                    document.add_paragraph(item, style="List Bullet")

        # ── 섹션 7: 기술 전망 ─────────────────────────────────────────────
        hc_phase = s7.get("hype_cycle_phase", "")
        predictions = s7.get("predictions", [])
        roadmap = s7.get("roadmap", [])

        has_hc = _has_real_text(hc_phase)
        has_predictions = _has_real_data(predictions, ["prediction"])
        has_roadmap = _has_real_data(roadmap, ["milestone"])

        if any([has_hc, has_predictions, has_roadmap]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"섹션 {sn} – 기술 전망", level=1)
            _ssn = [0]

            if has_hc:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 가트너 하이프 사이클 위치", level=2)
                hc = document.add_table(rows=3, cols=2)
                hc.style = "Light Grid Accent 1"
                _bold_cell(hc.rows[0].cells[0], "현재 단계")
                hc.rows[0].cells[1].text = hc_phase
                _bold_cell(hc.rows[1].cells[0], "고원 도달 예상 기간")
                hc.rows[1].cells[1].text = s7.get("hype_cycle_plateau", "N/A")
                _bold_cell(hc.rows[2].cells[0], "출처")
                hc.rows[2].cells[1].text = s7.get("hype_cycle_source", "N/A")
                document.add_paragraph()

            if has_predictions:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} IDC 방식 5년 예측", level=2)
                _add_table(
                    document,
                    ["연도", "예측", "신뢰도", "시사점", "출처"],
                    [
                        [r.get("year",""), r.get("prediction",""), r.get("confidence",""), r.get("implication",""), r.get("source","")]
                        for r in predictions
                    ],
                )

            if has_roadmap:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 기술 로드맵", level=2)
                _add_table(
                    document,
                    ["단계", "기간", "마일스톤", "TRL", "핵심 촉진 요인", "리스크"],
                    [
                        [r.get("phase",""), r.get("timeframe",""), r.get("milestone",""), r.get("trl",""), r.get("enabler",""), r.get("risk","")]
                        for r in roadmap
                    ],
                )

        # ── 섹션 8: 프라운호퍼 한국에 대한 전략적 시사점 ──────────────────
        opportunities = s8.get("opportunities", [])
        risks = s8.get("risks", [])
        actions = [a for a in s8.get("actions", []) if not _is_na(a)]

        has_opps = _has_real_data(opportunities, ["opportunity"])
        has_risks = _has_real_data(risks, ["risk"])
        has_actions = bool(actions)

        if any([has_opps, has_risks, has_actions]):
            _sn[0] += 1
            sn = _sn[0]
            document.add_heading(f"섹션 {sn} – 프라운호퍼 한국에 대한 전략적 시사점", level=1)
            _ssn = [0]

            if has_opps:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 기회 평가", level=2)
                _add_table(
                    document,
                    ["기회", "유형", "잠재적 파트너", "재원", "우선순위", "타임라인"],
                    [
                        [o.get("opportunity",""), o.get("type",""), o.get("partners",""), o.get("funding",""), o.get("priority",""), o.get("timeline","")]
                        for o in opportunities
                    ],
                )

            if has_risks:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 리스크 등록부", level=2)
                _add_table(
                    document,
                    ["리스크", "발생 가능성", "영향", "완화 방안", "담당자"],
                    [
                        [r.get("risk",""), r.get("likelihood",""), r.get("impact",""), r.get("mitigation",""), r.get("owner","")]
                        for r in risks
                    ],
                )

            if has_actions:
                _ssn[0] += 1
                document.add_heading(f"{sn}.{_ssn[0]} 권고 조치 사항 (향후 90일)", level=2)
                for action in actions:
                    document.add_paragraph(action, style="List Bullet")

        # ── 섹션 9: 출처 및 방법론 ────────────────────────────────────────
        _sn[0] += 1
        sn = _sn[0]
        document.add_heading(f"섹션 {sn} – 출처 및 방법론", level=1)
        _ssn = [0]

        _ssn[0] += 1
        document.add_heading(f"{sn}.{_ssn[0]} 출처 등록부", level=2)
        document.add_paragraph("학술 및 연구:").bold = True
        for src in ["IEEE Xplore – ieee.org/xplore", "ACM Digital Library – dl.acm.org",
                    "arXiv – arxiv.org", "Fraunhofer-Publica – publica.fraunhofer.de",
                    "ETRI Journal – etri.re.kr", "KISTI – kisti.re.kr",
                    "Springer / Elsevier / Wiley – 기관 구독을 통한 접근"]:
            document.add_paragraph(src, style="List Bullet")

        document.add_paragraph("시장 인텔리전스:").bold = True
        for src in ["Gartner – gartner.com", "IDC – idc.com", "Statista – statista.com",
                    "McKinsey Global Institute – mckinsey.com/mgi", "KIET – kiet.re.kr",
                    "KOTRA – kotra.or.kr", "KISDI – kisdi.re.kr", "ADB – adb.org",
                    "OECD iLibrary – oecd-ilibrary.org", "Eurostat – ec.europa.eu/eurostat",
                    "Destatis – destatis.de", "KOSTAT – kostat.go.kr"]:
            document.add_paragraph(src, style="List Bullet")

        document.add_paragraph("특허 및 IP:").bold = True
        for src in ["Espacenet (EPO) – espacenet.epo.org", "KIPO – kipo.go.kr",
                    "DPMA – dpma.de", "JPO – j-platpat.inpit.go.jp",
                    "CNIPA – cnipa.gov.cn", "Google Patents – patents.google.com"]:
            document.add_paragraph(src, style="List Bullet")

        document.add_paragraph("정부 및 정책:").bold = True
        for src in ["EU Commission – ec.europa.eu", "BMBF – bmbf.de",
                    "MOTIE – motie.go.kr", "MSIT – msit.go.kr",
                    "KISTEP – kistep.re.kr", "IITP – iitp.kr",
                    "NIST – nist.gov", "IEA – iea.org", "NEDO – nedo.go.jp",
                    "METI – meti.go.jp", "APEC – apec.org",
                    "NITI Aayog – niti.gov.in", "EDB Singapore – edb.gov.sg"]:
            document.add_paragraph(src, style="List Bullet")

        methodology = s9.get("methodology", "")
        if _has_real_text(methodology):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} 자동화 모니터링 방법론", level=2)
            document.add_paragraph(methodology)

        quality = s9.get("quality", "")
        if _has_real_text(quality):
            _ssn[0] += 1
            document.add_heading(f"{sn}.{_ssn[0]} 데이터 품질 및 한계", level=2)
            document.add_paragraph(quality)

        # ── 참고문헌 ──────────────────────────────────────────────────────
        document.add_page_break()
        document.add_heading("참고문헌 (References)", level=1)
        accessed = date.today().isoformat()
        _add_table(
            document,
            ["번호", "출처", "제목", "발행일", "URL"],
            [
                [
                    f"[{i}]",
                    item.get("source_name", item.get("category", "Unknown")),
                    item["title"],
                    item["log_date"],
                    item["url"],
                ]
                for i, item in enumerate(logs, start=1)
            ],
        )

        # ── 출처 색인 ─────────────────────────────────────────────────────
        document.add_heading("출처 색인 (Source Index)", level=1)
        src_tbl = document.add_table(rows=1, cols=4)
        src_tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["날짜", "제목", "카테고리", "URL"]):
            _bold_cell(src_tbl.rows[0].cells[i], h)
        for item in logs:
            row = src_tbl.add_row().cells
            row[0].text = item["log_date"]
            row[1].text = item["title"]
            row[2].text = CATEGORY_LABELS_KO.get(item["category"], item["category"])
            _add_hyperlink(row[3].paragraphs[0], item["url"], item["url"])

        document.add_paragraph(f"\n— 보고서 끝 — 생성일: {accessed}")
        return document

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_monthly_report(
        self,
        year: int,
        month: int,
        logs: list[dict],
    ) -> Path:
        """Generate an English TMR Word document following the Fraunhofer template."""
        structured = self._synthesize(year, month, logs, lang="en")
        document = self._build_document(year, month, logs, structured)

        output_path = self.output_dir / f"tech-market-report-{year}-{month:02d}.docx"
        document.save(output_path)
        logger.info("Generated Word report: %s", output_path)
        return output_path

    def generate_monthly_report_ko(
        self,
        year: int,
        month: int,
        logs: list[dict],
    ) -> Path:
        """Generate a Korean TMR Word document following the Fraunhofer template."""
        structured = self._synthesize(year, month, logs, lang="ko")
        document = self._build_document_ko(year, month, logs, structured)

        output_path = self.output_dir / f"tech-market-report-{year}-{month:02d}-ko.docx"
        document.save(output_path)
        logger.info("Generated Korean Word report: %s", output_path)
        return output_path
